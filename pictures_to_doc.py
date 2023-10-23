import os
from docx import Document
from docx.shared import Inches
from PIL import Image

# Funkce pro vytvoření DOCX dokumentu a vložení obrázku s popiskem
def create_docx_with_images(image_dir, docx_filename):
    # Vytvoření nového Word dokumentu
    doc = Document()
    doc.add_heading('Obrázky a popisky', 0)

    # Procházení adresářů a podadresářů
    for root, dirs, files in os.walk(image_dir):
        for file in files:
            # Získání cesty k souboru a jeho název
            file_path = os.path.join(root, file)
            file_name = os.path.basename(file_path)

            # Získání názvu adresáře
            dir_name = os.path.basename(root)

            # Pokud soubor je obrázek, přidáme ho do dokumentu
            if file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                # Přidáme popisek s názvem adresáře
                doc.add_paragraph(f'Popisek: {dir_name}')

                # Vložíme obrázek do dokumentu
                doc.add_picture(file_path, width=Inches(4.0))

    # Uložení dokumentu do DOCX souboru
    doc.save(docx_filename)

    print(f'Dokument byl úspěšně vytvořen jako {docx_filename}')

# Zadejte cestu k adresáři, který chcete procházet
image_directory = 'path_to_your_image_directory'

# Název DOCX souboru, kam budou obrázky uloženy
output_docx_filename = 'output_document.docx'

# Volání funkce pro vytvoření dokumentu
create_docx_with_images(image_directory, output_docx_filename)